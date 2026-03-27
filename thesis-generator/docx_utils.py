from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
import re
import os

def set_font_style(run, font_name_cn='宋体', font_name_en='Times New Roman', size=12, bold=False, superscript=False):
    """
    设置中西文混合字体样式
    size: 磅值 (pt)，例如小四对应 12
    """
    run.font.name = font_name_en
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name_cn)
    run.font.size = Pt(size)
    run.font.bold = bold
    if superscript:
        run.font.superscript = True

def format_paragraph(p, align=WD_PARAGRAPH_ALIGNMENT.JUSTIFY, line_spacing=1.5, space_before=0, space_after=0, first_indent=0):
    """
    设置段落格式
    space_before/after: 磅值 (pt)
    first_indent: 磅值 (pt)，例如首行缩进2字符约24pt (小四)
    """
    p.alignment = align
    p.paragraph_format.line_spacing = line_spacing
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.first_line_indent = Pt(first_indent)

def add_heading(doc, text, level):
    p = doc.add_paragraph()
    run = p.add_run(text)
    if level == 1: 
        set_font_style(run, size=16, bold=True)
        format_paragraph(p, align=WD_PARAGRAPH_ALIGNMENT.LEFT, space_before=24, space_after=24, first_indent=0)
    elif level == 2: 
        set_font_style(run, size=15, bold=True)
        format_paragraph(p, align=WD_PARAGRAPH_ALIGNMENT.LEFT, space_before=0, space_after=0, first_indent=0)
    elif level == 3: 
        set_font_style(run, size=14, bold=True)
        format_paragraph(p, align=WD_PARAGRAPH_ALIGNMENT.LEFT, space_before=0, space_after=0, first_indent=0)

def add_text_with_citations(doc, text):
    """
    添加包含 [1] 引用的正文
    """
    if not text: return
    p = doc.add_paragraph()
    format_paragraph(p, first_indent=24)
    parts = re.split(r'(\[\d+(?:-\d+)?\])', text)
    for part in parts:
        run = p.add_run(part)
        is_cit = re.match(r'\[\d+(?:-\d+)?\]', part)
        set_font_style(run, size=12, superscript=bool(is_cit))

def add_image(doc, img_path, caption, width_cm=12, force_page_break=False):
    """
    插入图片，支持自定义宽度和强制分页
    """
    if os.path.exists(img_path):
        if force_page_break:
            doc.add_page_break()
            
        p = doc.add_paragraph()
        format_paragraph(p, align=WD_PARAGRAPH_ALIGNMENT.CENTER, first_indent=0)
        run = p.add_run()
        run.add_picture(img_path, width=Cm(width_cm))
        
        p_cap = doc.add_paragraph()
        format_paragraph(p_cap, align=WD_PARAGRAPH_ALIGNMENT.CENTER, first_indent=0, space_before=6)
        run_cap = p_cap.add_run(caption)
        set_font_style(run_cap, size=10.5) 
        doc.add_paragraph() 
    else:
        print(f"Missing: {img_path}")

def set_cell_border(cell, **kwargs):
    """
    设置单元格边框，用于制作三线表
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)
    
    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def set_table_style_three_line(table):
    """
    快捷设置三线表样式
    """
    for row in table.rows:
        for cell in row.cells:
            # 默认无边框
            set_cell_border(cell, 
                left={"val": "nil"}, right={"val": "nil"}, 
                top={"val": "nil"}, bottom={"val": "nil"},
                insideH={"val": "nil"}, insideV={"val": "nil"}
            )
    
    # 顶线 (粗)
    for cell in table.rows[0].cells:
        set_cell_border(cell, top={"sz": 24, "val": "single", "space": "0"})
        # 标题行底线 (细)
        set_cell_border(cell, bottom={"sz": 4, "val": "single", "space": "0"})
        
    # 底线 (粗)
    for cell in table.rows[-1].cells:
        set_cell_border(cell, bottom={"sz": 24, "val": "single", "space": "0"})

def add_custom_header_footer(section, header_text):
    """
    添加带双线边框的页眉和标准页脚
    """
    # Header
    header = section.header
    header.is_linked_to_previous = False
    p = header.paragraphs[0]
    p.clear()
    
    # 0.5磅双线底边框 XML
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'double')
    bottom.set(qn('w:sz'), '4') # 4 = 1/2 pt
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'auto')
    pBdr.append(bottom)
    p._p.get_or_add_pPr().append(pBdr)
    
    run = p.add_run(header_text)
    set_font_style(run, size=9) # 小五
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    r1 = p.add_run("第 ")
    set_font_style(r1, size=9)
    # Page field
    fld1 = OxmlElement('w:fldChar'); fld1.set(qn('w:fldCharType'), 'begin')
    instr = OxmlElement('w:instrText'); instr.set(qn('xml:space'), 'preserve'); instr.text = "PAGE"
    fld2 = OxmlElement('w:fldChar'); fld2.set(qn('w:fldCharType'), 'end')
    r1._element.append(fld1); r1._element.append(instr); r1._element.append(fld2)
    
    r2 = p.add_run(" 页 共 ")
    set_font_style(r2, size=9)
    
    # NumPages field
    fld3 = OxmlElement('w:fldChar'); fld3.set(qn('w:fldCharType'), 'begin')
    instr2 = OxmlElement('w:instrText'); instr2.set(qn('xml:space'), 'preserve'); instr2.text = "NUMPAGES"
    fld4 = OxmlElement('w:fldChar'); fld4.set(qn('w:fldCharType'), 'end')
    r2._element.append(fld3); r2._element.append(instr2); r2._element.append(fld4)
    
    r3 = p.add_run(" 页")
    set_font_style(r3, size=9)
